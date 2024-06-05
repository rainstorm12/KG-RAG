import docx
import os
def list_filenames(folder_path):
    """
    列出指定文件夹路径下的所有文件名。
    
    参数:
    folder_path -- 目标文件夹的路径字符串
    
    返回值:
    一个包含所有文件名的列表
    """
    # 确保传入的是一个字符串
    if not isinstance(folder_path, str):
        raise ValueError("folder_path 必须是字符串类型")
    
    # 初始化一个空列表来存储文件名
    filenames = []
    
    # 使用os.walk()函数遍历文件夹
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            # 构建完整的文件路径
            full_path = os.path.join(root, file)
            # 将文件名添加到列表中
            filenames.append(os.path.basename(full_path))
    
    return filenames

if __name__ == '__main__':
    folder_path = "./data/pipe/"
    txt_path = './data/pipe.txt'

    filenames = list_filenames(folder_path)
    doc_text = ""
    for filename in filenames:
        doc = docx.Document(folder_path + filename)
        for i in range(0,len(doc.paragraphs)):
            for j in range(0,len(doc.paragraphs[i].runs)):
                doc_text += doc.paragraphs[i].runs[j].text
                
    with open(txt_path,'w') as f:
        f.write(doc_text)